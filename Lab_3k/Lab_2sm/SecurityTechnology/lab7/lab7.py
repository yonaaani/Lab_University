# Використовую бібліотеку Crypto від PyCrypto та PyCryptodome
from Crypto.PublicKey import RSA
from Crypto.Cipher import PKCS1_OAEP
# Ще одна бібліотека для використання Microsoft Word
from docx import Document
import base64

# Ця функція генерує пару асиметричних ключів RSA довжиною 2048 біт
def generate_key_pair():
    key = RSA.generate(2048)
    private_key = key.export_key()
    public_key = key.publickey().export_key()
    return private_key, public_key

# Ця функція приймає ім'я файлу filename та ключ key, і записує цей ключ у файл у бінарному(двійковому) режимі
def save_key(filename, key):
    with open(filename, 'wb') as f:
        f.write(key)

# Ця функція приймає ім'я файлу filename і читає з нього ключ, повертаючи його у вигляді байтового рядка
def load_key(filename):
    with open(filename, 'rb') as f:
        key = f.read()
    return key

def encrypt_file(input_file, output_file, public_key):
    # Зчитування документу Microsoft Word
    doc = Document(input_file)
    plaintext = ""
    for paragraph in doc.paragraphs:
        plaintext += paragraph.text + "\n"
    # Шифрування тексту за допомогою асиметричного шифрування RSA
    recipient_key = RSA.import_key(public_key)
    cipher_rsa = PKCS1_OAEP.new(recipient_key)
    encrypted_data = cipher_rsa.encrypt(plaintext.encode())
    # Кодування зашифрованих даних у формат base64
    encoded_data = base64.b64encode(encrypted_data).decode('utf-8')
    # Запис зашифрованих даних у документ Microsoft Word
    encrypted_doc = Document()
    encrypted_doc.add_paragraph(encoded_data)
    encrypted_doc.save(output_file)

def decrypt_file(input_file, output_file, private_key):
    # Зчитування зашифрованого документу Microsoft Word
    encrypted_doc = Document(input_file)
    encoded_data = ""
    for paragraph in encrypted_doc.paragraphs:
        encoded_data += paragraph.text
    # Декодування зашифрованих даних з формату base64
    encrypted_data = base64.b64decode(encoded_data.encode('utf-8'))
    # Розшифрування тексту за допомогою асиметричного шифрування RSA
    private_key = RSA.import_key(private_key)
    cipher_rsa = PKCS1_OAEP.new(private_key)
    decrypted_data = cipher_rsa.decrypt(encrypted_data).decode('utf-8')
    # Запис розшифрованого тексту у документ Microsoft Word
    decrypted_doc = Document()
    decrypted_doc.add_paragraph(decrypted_data)
    decrypted_doc.save(output_file)
    
# Згенерувати ключі
private_key, public_key = generate_key_pair()
save_key("private.pem", private_key)
save_key("public.pem", public_key)

print("Private Key:\n", private_key)
print("Public Key:\n", public_key)

# Шифрування файлу
encrypt_file("C:/Users/yonaaani/OneDrive/Рабочий стол/lab7/plaintext.docx", "C:/Users/yonaaani/OneDrive/Рабочий стол/lab7/encrypted.docx", public_key)

# Розшифрування файлу
decrypt_file("C:/Users/yonaaani/OneDrive/Рабочий стол/lab7/encrypted.docx", "C:/Users/yonaaani/OneDrive/Рабочий стол/lab7/decrypted.docx", private_key)

print("Sucess!")